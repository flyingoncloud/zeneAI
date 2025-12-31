#!/usr/bin/env python3
"""
DOCX File Inspector
Inspects DOCX files to show content and check for placeholders.
"""

import os
from docx import Document

def inspect_docx(file_path):
    """Inspect a DOCX file and show its content."""
    
    if not os.path.exists(file_path):
        print(f"❌ File not found: {file_path}")
        return
    
    try:
        doc = Document(file_path)
        
        print(f"📄 Inspecting: {file_path}")
        print(f"📊 File size: {os.path.getsize(file_path)} bytes")
        print("=" * 60)
        
        # Show paragraphs
        print("\n📝 DOCUMENT CONTENT (first 30 paragraphs):")
        print("-" * 40)
        
        paragraph_count = 0
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                paragraph_count += 1
                if paragraph_count <= 30:
                    print(f"{paragraph_count:2d}: {para.text}")
        
        print(f"\nTotal paragraphs with content: {paragraph_count}")
        
        # Check for placeholders
        print("\n🔍 CHECKING FOR PLACEHOLDERS:")
        print("-" * 40)
        
        all_text = '\n'.join([p.text for p in doc.paragraphs])
        
        placeholders = [
            '{{REPORT_DATE}}',
            '{{CONVERSATION_ID}}', 
            '{{TOTAL_MESSAGES}}',
            '{{EXECUTIVE_SUMMARY}}',
            '{{CONVERSATION_OVERVIEW}}',
            '{{FRAMEWORK_ANALYSIS}}',
            '{{THERAPEUTIC_INSIGHTS}}',
            '{{RECOMMENDATIONS}}'
        ]
        
        found_placeholders = []
        for placeholder in placeholders:
            if placeholder in all_text:
                print(f"✅ Found: {placeholder}")
                found_placeholders.append(placeholder)
            else:
                print(f"❌ Missing: {placeholder}")
        
        print(f"\nPlaceholders found: {len(found_placeholders)}/{len(placeholders)}")
        
        # Check for Chinese content
        print("\n🇨🇳 CHINESE CONTENT CHECK:")
        print("-" * 40)
        
        chinese_chars = 0
        for char in all_text:
            if '\u4e00' <= char <= '\u9fff':  # Chinese character range
                chinese_chars += 1
        
        print(f"Chinese characters: {chinese_chars}")
        print(f"Total characters: {len(all_text)}")
        
        if chinese_chars > 0:
            print(f"Chinese content: {chinese_chars/len(all_text)*100:.1f}%")
        
        # Show tables if any
        if doc.tables:
            print(f"\n📊 TABLES FOUND: {len(doc.tables)}")
            for i, table in enumerate(doc.tables[:3]):
                print(f"Table {i+1}: {len(table.rows)} rows, {len(table.columns)} columns")
        
        return True
        
    except Exception as e:
        print(f"❌ Error reading DOCX: {e}")
        return False

if __name__ == "__main__":
    # Inspect both template files
    files_to_inspect = [
        "src/resources/ZENE_Report_Pro_Edited_25Dec2025.docx",
        "src/resources/ZENE_Chinese_Template.docx"
    ]
    
    for file_path in files_to_inspect:
        inspect_docx(file_path)
        print("\n" + "="*80 + "\n")