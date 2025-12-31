#!/usr/bin/env python3
"""
Update Chinese Template with Psychological Scores

This script adds the psychological scoring section to the Chinese template.
"""

import os
from docx import Document

def update_chinese_template():
    """Update the Chinese template to include psychological scores placeholder"""
    
    template_path = "src/resources/ZENE_Chinese_Template.docx"
    
    if not os.path.exists(template_path):
        print(f"❌ Template not found: {template_path}")
        return False
    
    try:
        # Open the template
        doc = Document(template_path)
        
        # Find the position to insert psychological scores (after title, before executive summary)
        target_paragraph = None
        for i, para in enumerate(doc.paragraphs):
            if "执行摘要" in para.text:
                target_paragraph = i
                break
        
        if target_paragraph is None:
            print("❌ Could not find insertion point in template")
            return False
        
        # Insert psychological scores section before executive summary
        psychological_scores_para = doc.paragraphs[target_paragraph].insert_paragraph_before()
        psychological_scores_para.text = "五大核心心智指数"
        
        scores_placeholder_para = doc.paragraphs[target_paragraph + 1].insert_paragraph_before()
        scores_placeholder_para.text = "{{PSYCHOLOGICAL_SCORES}}"
        
        # Add empty line
        empty_para = doc.paragraphs[target_paragraph + 2].insert_paragraph_before()
        empty_para.text = ""
        
        # Save the updated template
        doc.save(template_path)
        print(f"✅ Updated Chinese template with psychological scores placeholder")
        return True
        
    except Exception as e:
        print(f"❌ Error updating template: {e}")
        return False

if __name__ == "__main__":
    success = update_chinese_template()
    if success:
        print("🎉 Chinese template updated successfully!")
    else:
        print("❌ Failed to update Chinese template")