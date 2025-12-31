#!/usr/bin/env python3
"""
Create Chinese Template

Creates a Chinese DOCX template for psychology reports.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_chinese_template():
    """Create a Chinese template for psychology reports"""
    
    # Create new document
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title
    title = doc.add_paragraph("ZENE心理分析报告", style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph("多框架治疗评估", style='Normal')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Space
    
    # Metadata table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # Report metadata with placeholders
    metadata = [
        ("报告日期", "{{REPORT_DATE}}"),
        ("分析系统", "ZENE多框架心理检测系统"),
        ("对话ID", "{{CONVERSATION_ID}}"),
        ("总消息数", "{{TOTAL_MESSAGES}}"),
        ("分析版本", "v2.0 - 多框架集成")
    ]
    
    for i, (key, value) in enumerate(metadata):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = value
    
    # Page break
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading("执行摘要", level=1)
    doc.add_paragraph("{{EXECUTIVE_SUMMARY}}")
    doc.add_paragraph()
    
    # Conversation Overview
    doc.add_heading("对话概览", level=1)
    doc.add_paragraph("{{CONVERSATION_OVERVIEW}}")
    doc.add_paragraph()
    
    # Framework Analysis
    doc.add_heading("心理框架分析", level=1)
    doc.add_paragraph("{{FRAMEWORK_ANALYSIS}}")
    doc.add_paragraph()
    
    # Therapeutic Insights
    doc.add_heading("治疗洞察", level=1)
    doc.add_paragraph("{{THERAPEUTIC_INSIGHTS}}")
    doc.add_paragraph()
    
    # Recommendations
    doc.add_heading("治疗建议", level=1)
    doc.add_paragraph("{{RECOMMENDATIONS}}")
    doc.add_paragraph()
    
    # Technical Appendix
    doc.add_heading("技术附录", level=1)
    
    # Analysis methodology
    method_para = doc.add_paragraph()
    method_para.add_run("分析方法：").bold = True
    method_para.add_run("\\n本报告使用ZENE多框架心理检测系统生成，采用两阶段混合方法：\\n")
    method_para.add_run("1. 模式匹配：使用语言和语义模式进行初步筛选\\n")
    method_para.add_run("2. LLM分析：使用大型语言模型进行心理要素深度检测\\n")
    method_para.add_run("3. 置信度评分：为每个检测要素提供统计置信度测量\\n")
    method_para.add_run("4. 跨框架整合：分析多个治疗框架之间的相互作用\\n")
    
    doc.add_paragraph()
    
    # Framework details
    frameworks_para = doc.add_paragraph()
    frameworks_para.add_run("支持的框架：").bold = True
    frameworks_para.add_run("\\n• IFS（内在家庭系统）：部分工作和自我领导\\n")
    frameworks_para.add_run("• CBT（认知行为疗法）：认知扭曲和行为模式\\n")
    frameworks_para.add_run("• 荣格心理学：原型、梦境和个体化\\n")
    frameworks_para.add_run("• 叙事疗法：外化和故事重新创作\\n")
    frameworks_para.add_run("• 依恋理论：关系模式和情绪调节\\n")
    
    doc.add_paragraph()
    
    # Disclaimer
    disclaimer_para = doc.add_paragraph()
    disclaimer_para.add_run("免责声明：").bold = True
    disclaimer_para.add_run("本报告由AI系统生成，仅供参考。不应替代专业心理评估或临床判断。所有治疗决策应在合格心理健康专业人员的指导下做出。")
    
    # Save template
    template_dir = "src/resources"
    os.makedirs(template_dir, exist_ok=True)
    template_path = os.path.join(template_dir, "ZENE_Chinese_Template.docx")
    doc.save(template_path)
    
    print(f"✅ 中文模板已创建: {template_path}")
    return template_path

if __name__ == "__main__":
    create_chinese_template()