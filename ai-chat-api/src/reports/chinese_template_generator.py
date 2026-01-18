"""
ZENE Chinese Template-Based Report Generator

Generates comprehensive psychology analysis reports in Chinese using DOCX templates
when conversations contain sufficient psychological information.
"""

import os
import logging
import shutil
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class ChineseTemplateReportGenerator:
    """
    Generates professional psychology analysis reports in Chinese using DOCX templates.
    
    Creates comprehensive reports including:
    - 执行摘要 (Executive summary)
    - 框架分析 (Framework analysis)  
    - 治疗建议 (Therapeutic recommendations)
    - 对话洞察 (Conversation insights)
    """
    
    def __init__(self, template_path: str = None, output_dir: str = "reports"):
        # Use the correct Chinese template with placeholders
        self.template_path = template_path or "src/resources/ZENE_Chinese_Template.docx"
        # Fallback to the other template only if the Chinese template doesn't exist
        if not os.path.exists(self.template_path):
            self.template_path = "src/resources/ZENE_Report_Pro_Edited_25Dec2025.docx"
        
        self.output_dir = output_dir
        self.ensure_output_directory()
    
    def ensure_output_directory(self):
        """Ensure the output directory exists"""
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def should_generate_report(self, conversation_data: Dict) -> Tuple[bool, str]:
        """
        Determine if a conversation has enough information to generate a report.
        
        Args:
            conversation_data: Complete conversation data with psychology analysis
            
        Returns:
            Tuple of (should_generate, reason)
        """
        messages = conversation_data.get('messages', [])
        psychology_analyses = []
        
        # Collect all psychology analyses from messages
        for message in messages:
            if message.get('psychology_analysis', {}).get('analyzed', False):
                psychology_analyses.append(message['psychology_analysis'])
        
        if not psychology_analyses:
            return False, "对话中未找到心理分析数据"
        
        # Check for minimum conversation length
        if len(messages) < 6:
            return False, f"对话过短（{len(messages)}条消息，至少需要6条）"
        
        # Check for framework diversity
        all_frameworks = set()
        total_confidence = 0
        analysis_count = 0
        
        for analysis in psychology_analyses:
            frameworks = analysis.get('frameworks', {})
            for framework_name, framework_data in frameworks.items():
                confidence = framework_data.get('confidence_score', 0.0)
                elements = framework_data.get('elements_detected', [])
                
                if confidence >= 0.5 or len(elements) >= 2:
                    all_frameworks.add(framework_name)
                    total_confidence += confidence
                    analysis_count += 1
        
        # Require at least 2 frameworks detected
        if len(all_frameworks) < 2:
            return False, f"框架多样性不足（{len(all_frameworks)}个框架，至少需要2个）"
        
        # Require minimum average confidence
        avg_confidence = total_confidence / max(analysis_count, 1)
        if avg_confidence < 0.6:
            return False, f"置信度分数过低（平均{avg_confidence:.2f}，至少需要0.6）"
        
        return True, f"报告标准已满足：{len(all_frameworks)}个框架，{avg_confidence:.2f}平均置信度"
    
    def generate_report(self, conversation_data: Dict, user_info: Optional[Dict] = None) -> str:
        """
        Generate a comprehensive psychology report using Chinese template.
        
        Args:
            conversation_data: Complete conversation data
            user_info: Optional user information
            
        Returns:
            Path to generated report file
        """
        # Generate filename with the exact format you requested
        timestamp = datetime.now().strftime("%d%b%Y")
        filename = f"ZENE_Report_Pro_Edited_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        # Copy template to output location
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"模板文件未找到: {self.template_path}")
        
        shutil.copy2(self.template_path, filepath)
        
        # Open the copied template
        doc = Document(filepath)
        
        # Generate Chinese report content
        self._update_template_content(doc, conversation_data, user_info)
        
        # Save document
        doc.save(filepath)
        logger.info(f"生成心理分析报告: {filepath}")
        
        return filepath
    
    def _update_template_content(self, doc: Document, conversation_data: Dict, user_info: Optional[Dict]):
        """Update template with Chinese content"""
        
        # Replace placeholders in the template
        self._replace_text_in_document(doc, "{{REPORT_DATE}}", datetime.now().strftime("%Y年%m月%d日"))
        self._replace_text_in_document(doc, "{{CONVERSATION_ID}}", str(conversation_data.get('id', 'N/A')))
        self._replace_text_in_document(doc, "{{TOTAL_MESSAGES}}", str(len(conversation_data.get('messages', []))))
        
        # Add psychological scoring system
        psychological_scores = self._generate_psychological_scores(conversation_data)
        self._replace_text_in_document(doc, "{{PSYCHOLOGICAL_SCORES}}", psychological_scores)
        
        # Add executive summary
        executive_summary = self._generate_chinese_executive_summary(conversation_data)
        self._replace_text_in_document(doc, "{{EXECUTIVE_SUMMARY}}", executive_summary)
        
        # Add conversation overview
        conversation_overview = self._generate_chinese_conversation_overview(conversation_data)
        self._replace_text_in_document(doc, "{{CONVERSATION_OVERVIEW}}", conversation_overview)
        
        # Add framework analysis
        framework_analysis = self._generate_chinese_framework_analysis(conversation_data)
        self._replace_text_in_document(doc, "{{FRAMEWORK_ANALYSIS}}", framework_analysis)
        
        # Add therapeutic insights
        therapeutic_insights = self._generate_chinese_therapeutic_insights(conversation_data)
        self._replace_text_in_document(doc, "{{THERAPEUTIC_INSIGHTS}}", therapeutic_insights)
        
        # Add recommendations
        recommendations = self._generate_chinese_recommendations(conversation_data)
        self._replace_text_in_document(doc, "{{RECOMMENDATIONS}}", recommendations)
    
    def _replace_text_in_document(self, doc: Document, placeholder: str, replacement: str):
        """Replace placeholder text in document"""
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, replacement)
    
    def _generate_chinese_executive_summary(self, conversation_data: Dict) -> str:
        """Generate executive summary in Chinese"""
        messages = conversation_data.get('messages', [])
        psychology_analyses = [msg.get('psychology_analysis') for msg in messages 
                             if msg.get('psychology_analysis', {}).get('analyzed', False)]
        
        if not psychology_analyses:
            return "本次对话分析中未发现心理分析数据。"
        
        # Framework summary
        all_frameworks = {}
        for analysis in psychology_analyses:
            frameworks = analysis.get('frameworks', {})
            for name, data in frameworks.items():
                if name not in all_frameworks:
                    all_frameworks[name] = {
                        'detections': 0,
                        'total_confidence': 0,
                        'elements': []
                    }
                
                confidence = data.get('confidence_score', 0.0)
                elements = data.get('elements_detected', [])
                
                if confidence > 0.3 or elements:
                    all_frameworks[name]['detections'] += 1
                    all_frameworks[name]['total_confidence'] += confidence
                    all_frameworks[name]['elements'].extend(elements)
        
        # Framework name mapping to Chinese
        framework_names_chinese = {
            'cbt': '认知行为疗法(CBT)',
            'ifs': '内在家庭系统(IFS)',
            'jungian': '荣格心理学',
            'narrative': '叙事疗法',
            'attachment': '依恋理论'
        }
        
        # Generate summary text
        summary_parts = []
        
        if all_frameworks:
            detected_frameworks = [name for name, data in all_frameworks.items() 
                                 if data['detections'] > 0]
            chinese_frameworks = [framework_names_chinese.get(name, name.upper()) for name in detected_frameworks]
            summary_parts.append(f"本次对话分析检测到来自{len(detected_frameworks)}个心理学框架的模式：{', '.join(chinese_frameworks)}。")
        
        # Primary framework
        if all_frameworks:
            primary_framework = max(all_frameworks.items(), 
                                  key=lambda x: x[1]['total_confidence'] / max(x[1]['detections'], 1))
            avg_confidence = primary_framework[1]['total_confidence'] / max(primary_framework[1]['detections'], 1)
            chinese_name = framework_names_chinese.get(primary_framework[0], primary_framework[0].upper())
            summary_parts.append(f"主要识别的治疗框架为{chinese_name}，平均置信度为{avg_confidence:.2f}。")
        
        # Conversation characteristics
        user_messages = [msg for msg in messages if msg.get('role') == 'user']
        if user_messages:
            summary_parts.append(f"对话包含{len(user_messages)}条用户消息，涵盖多个治疗领域。")
        
        summary_parts.append("本报告提供了检测到的心理模式的详细分析、治疗洞察和持续支持建议。")
        
        return "\\n\\n".join(summary_parts)
    
    def _generate_chinese_conversation_overview(self, conversation_data: Dict) -> str:
        """Generate conversation overview in Chinese"""
        messages = conversation_data.get('messages', [])
        user_messages = [msg for msg in messages if msg.get('role') == 'user']
        ai_messages = [msg for msg in messages if msg.get('role') == 'assistant']
        
        overview_parts = []
        overview_parts.append("对话统计信息：")
        overview_parts.append(f"• 总消息数：{len(messages)}")
        overview_parts.append(f"• 用户消息：{len(user_messages)}")
        overview_parts.append(f"• AI回复：{len(ai_messages)}")
        
        if messages:
            start_time = messages[0].get('timestamp', '未知')
            end_time = messages[-1].get('timestamp', '未知')
            overview_parts.append(f"• 对话时间：{start_time} 至 {end_time}")
        
        # Key themes (first few user messages)
        if user_messages:
            overview_parts.append("")
            overview_parts.append("讨论的主要主题：")
            
            for i, msg in enumerate(user_messages[:3], 1):
                content = msg.get('content', '')[:100]
                if len(msg.get('content', '')) > 100:
                    content += "..."
                overview_parts.append(f"{i}. {content}")
        
        return "\\n".join(overview_parts)
    
    def _generate_chinese_framework_analysis(self, conversation_data: Dict) -> str:
        """Generate detailed framework analysis in Chinese"""
        messages = conversation_data.get('messages', [])
        psychology_analyses = [msg.get('psychology_analysis') for msg in messages 
                             if msg.get('psychology_analysis', {}).get('analyzed', False)]
        
        if not psychology_analyses:
            return "无框架分析数据可用。"
        
        framework_names_chinese = {
            'cbt': '认知行为疗法(CBT)',
            'ifs': '内在家庭系统(IFS)',
            'jungian': '荣格心理学',
            'narrative': '叙事疗法',
            'attachment': '依恋理论'
        }
        
        # Aggregate framework data
        framework_summary = {}
        for analysis in psychology_analyses:
            frameworks = analysis.get('frameworks', {})
            for name, data in frameworks.items():
                if name not in framework_summary:
                    framework_summary[name] = {
                        'total_confidence': 0,
                        'detection_count': 0,
                        'all_elements': [],
                        'highest_confidence': 0
                    }
                
                confidence = data.get('confidence_score', 0.0)
                elements = data.get('elements_detected', [])
                
                if confidence > 0.3 or elements:
                    framework_summary[name]['total_confidence'] += confidence
                    framework_summary[name]['detection_count'] += 1
                    framework_summary[name]['all_elements'].extend(elements)
                    framework_summary[name]['highest_confidence'] = max(
                        framework_summary[name]['highest_confidence'], confidence
                    )
        
        # Sort frameworks by average confidence
        sorted_frameworks = sorted(
            framework_summary.items(),
            key=lambda x: x[1]['total_confidence'] / max(x[1]['detection_count'], 1),
            reverse=True
        )
        
        analysis_parts = []
        
        for framework_name, data in sorted_frameworks:
            if data['detection_count'] == 0:
                continue
                
            avg_confidence = data['total_confidence'] / data['detection_count']
            chinese_name = framework_names_chinese.get(framework_name, framework_name.upper())
            
            # Framework heading
            analysis_parts.append(f"{chinese_name}框架分析")
            analysis_parts.append(f"检测频率：{data['detection_count']}次")
            analysis_parts.append(f"平均置信度：{avg_confidence:.2f}")
            analysis_parts.append(f"峰值置信度：{data['highest_confidence']:.2f}")
            
            # Elements detected
            if data['all_elements']:
                analysis_parts.append("检测到的关键要素：")
                
                # Group elements by type
                element_types = {}
                for element in data['all_elements']:
                    elem_type = element.get('type', 'unknown')
                    if elem_type not in element_types:
                        element_types[elem_type] = []
                    element_types[elem_type].append(element)
                
                for elem_type, elements in element_types.items():
                    subtypes = [elem.get('subtype', elem_type) for elem in elements[:3]]
                    analysis_parts.append(f"• {elem_type}: {', '.join(set(subtypes))}")
            
            # Framework interpretation
            interpretation = self._get_chinese_framework_interpretation(framework_name, data)
            if interpretation:
                analysis_parts.append(f"临床解释：{interpretation}")
            
            analysis_parts.append("")  # Space between frameworks
        
        return "\\n".join(analysis_parts)
    
    def _generate_chinese_therapeutic_insights(self, conversation_data: Dict) -> str:
        """Generate therapeutic insights in Chinese"""
        messages = conversation_data.get('messages', [])
        psychology_analyses = [msg.get('psychology_analysis') for msg in messages 
                             if msg.get('psychology_analysis', {}).get('analyzed', False)]
        
        insights = []
        
        # Cross-framework insights
        for analysis in psychology_analyses:
            cross_insights = analysis.get('cross_framework_insights', {})
            if cross_insights.get('multiple_frameworks_detected'):
                frameworks = cross_insights['multiple_frameworks_detected'].get('frameworks', [])
                if len(frameworks) > 1:
                    insights.append(f"检测到跨{', '.join(frameworks)}框架的多模态表现，表明复杂的心理动态需要综合治疗方法。")
        
        # Framework-specific insights in Chinese
        framework_insights_chinese = {
            'ifs': "内在家庭系统模式表明活跃的内在部分需要自我领导和部分整合工作。",
            'cbt': "认知行为模式表明有机会进行认知重构和行为干预。",
            'jungian': "荣格元素表明丰富的象征内容和个体化过程值得探索。",
            'narrative': "叙事疗法模式表明外化和故事重新创作方法的潜力。",
            'attachment': "依恋模式表明关系动态可能受益于依恋导向的干预。"
        }
        
        # Add framework-specific insights based on detections
        detected_frameworks = set()
        for analysis in psychology_analyses:
            frameworks = analysis.get('frameworks', {})
            for name, data in frameworks.items():
                if data.get('confidence_score', 0) > 0.5 or data.get('elements_detected', []):
                    detected_frameworks.add(name)
        
        for framework in detected_frameworks:
            if framework in framework_insights_chinese:
                insights.append(framework_insights_chinese[framework])
        
        # Add insights to document
        if insights:
            insight_parts = []
            for i, insight in enumerate(insights, 1):
                insight_parts.append(f"{i}. {insight}")
            return "\\n".join(insight_parts)
        else:
            return "本次对话未产生特定的治疗洞察。"
    
    def _generate_chinese_recommendations(self, conversation_data: Dict) -> str:
        """Generate therapeutic recommendations in Chinese"""
        messages = conversation_data.get('messages', [])
        psychology_analyses = [msg.get('psychology_analysis') for msg in messages 
                             if msg.get('psychology_analysis', {}).get('analyzed', False)]
        
        # Analyze detected frameworks for recommendations
        detected_frameworks = {}
        for analysis in psychology_analyses:
            frameworks = analysis.get('frameworks', {})
            for name, data in frameworks.items():
                confidence = data.get('confidence_score', 0)
                elements = data.get('elements_detected', [])
                if confidence > 0.5 or len(elements) >= 2:
                    if name not in detected_frameworks:
                        detected_frameworks[name] = 0
                    detected_frameworks[name] += confidence
        
        # Generate recommendations based on strongest frameworks
        recommendations = []
        
        if 'ifs' in detected_frameworks:
            recommendations.append("考虑内在家庭系统(IFS)疗法来探索和整合内在部分，培养自我领导和内在和谐。")
        
        if 'cbt' in detected_frameworks:
            recommendations.append("认知行为疗法(CBT)干预可能有助于解决已识别的认知扭曲并发展更平衡的思维模式。")
        
        if 'attachment' in detected_frameworks:
            recommendations.append("依恋导向疗法可以解决关系模式并支持安全依恋策略的发展。")
        
        if 'narrative' in detected_frameworks:
            recommendations.append("叙事疗法方法可能有助于外化问题并支持重新创作首选的生活故事。")
        
        if 'jungian' in detected_frameworks:
            recommendations.append("荣格分析方法可以探索象征内容、原型模式并支持个体化过程。")
        
        # General recommendations
        if len(detected_frameworks) > 1:
            recommendations.append("鉴于多框架表现，从多种模式中汲取的综合治疗方法可能最有益。")
        
        recommendations.append("使用经过验证的心理测量工具定期评估和监测治疗进展。")
        recommendations.append("考虑协作治疗计划，确保干预与来访者目标和偏好保持一致。")
        
        # Add recommendations to document
        recommendation_parts = []
        for i, rec in enumerate(recommendations, 1):
            recommendation_parts.append(f"{i}. {rec}")
        
        return "\\n".join(recommendation_parts)
    
    def _generate_psychological_scores(self, conversation_data: Dict) -> str:
        """Generate psychological scoring system based on conversation analysis"""
        messages = conversation_data.get('messages', [])
        psychology_analyses = [msg.get('psychology_analysis') for msg in messages 
                             if msg.get('psychology_analysis', {}).get('analyzed', False)]
        
        if not psychology_analyses:
            # Default scores if no analysis available
            scores = {
                'emotional_regulation': 50,
                'cognitive_flexibility': 50,
                'relationship_sensitivity': 50,
                'internal_conflict': 50,
                'growth_potential': 50
            }
        else:
            scores = self._calculate_psychological_scores(psychology_analyses)
        
        # Generate the scoring section
        scoring_parts = []
        scoring_parts.append("第2部 五大核心心智指数详细分析")
        scoring_parts.append("")
        scoring_parts.append("本图展示 ZENE 心智结构评估中的五大关键指标：")
        scoring_parts.append("")
        
        # 1. Emotional Regulation Index
        scoring_parts.append(f"1. 情绪调节能力指数（{scores['emotional_regulation']}）")
        scoring_parts.append("评估情绪识别、表达与调节的稳定性和有效性。")
        scoring_parts.append("分数越高，表示越能在压力下保持情绪平衡。")
        scoring_parts.append("")
        
        # 2. Cognitive Flexibility Index
        scoring_parts.append(f"2. 认知灵活度指数（{scores['cognitive_flexibility']}）")
        scoring_parts.append("评估思维弹性、观点切换能力与面对变化时的适应性。")
        scoring_parts.append("分数越高，表示越容易从不同角度理解问题。")
        scoring_parts.append("")
        
        # 3. Relationship Sensitivity Index
        scoring_parts.append(f"3. 关系敏感度指数（{scores['relationship_sensitivity']}）")
        scoring_parts.append("评估人际关系中的情绪共鸣能力、边界感与反应模式。")
        scoring_parts.append("分数偏高表示共情强，但也可能对他人反应比较敏感。")
        scoring_parts.append("")
        
        # 4. Internal Conflict Index (reverse indicator)
        scoring_parts.append(f"4. 内在冲突度指数（{scores['internal_conflict']}）")
        scoring_parts.append("评估内心矛盾结构、自我一致性与心理整合水平。")
        scoring_parts.append("此项为反向指标：分数越高，内在冲突越强。")
        scoring_parts.append("")
        
        # 5. Growth Potential Index
        scoring_parts.append(f"5. 成长潜能指数（{scores['growth_potential']}）")
        scoring_parts.append("评估学习意愿、心理韧性、自我修复与成长动力。")
        scoring_parts.append("分数越高，表示未来可发展性越强。")
        scoring_parts.append("")
        
        # Detailed analysis sections
        scoring_parts.append("五大核心心智指数详细分析")
        scoring_parts.append("")
        
        # Emotional Insight Analysis
        scoring_parts.append("2.1 情绪觉察（Emotional Insight Analysis）")
        scoring_parts.append(f"你的情绪调节能力当前评估分值为：{scores['emotional_regulation']} / 100")
        scoring_parts.append("")
        scoring_parts.append("情绪识别、表达能力：")
        emotion_level = self._get_emotion_level(scores['emotional_regulation'])
        scoring_parts.append(f"{emotion_level['identification']}")
        scoring_parts.append("")
        scoring_parts.append("情绪调节和恢复能力：")
        scoring_parts.append(f"{emotion_level['regulation']}")
        scoring_parts.append("")
        scoring_parts.append("情绪倾向与风险指数：")
        scoring_parts.append(f"{emotion_level['risk']}")
        scoring_parts.append("")
        
        # Impact analysis
        impact_analysis = self._get_score_impact_analysis(scores['emotional_regulation'], 'emotional_regulation')
        scoring_parts.append("对用户的影响")
        scoring_parts.append(impact_analysis)
        scoring_parts.append("")
        
        return "\\n".join(scoring_parts)
    
    def _calculate_psychological_scores(self, psychology_analyses: List[Dict]) -> Dict[str, int]:
        """Calculate psychological scores based on framework analysis"""
        
        # Initialize base scores
        scores = {
            'emotional_regulation': 50,
            'cognitive_flexibility': 50,
            'relationship_sensitivity': 50,
            'internal_conflict': 50,
            'growth_potential': 50
        }
        
        # Analyze frameworks and adjust scores
        framework_impacts = {
            'cbt': {
                'emotional_regulation': 5,
                'cognitive_flexibility': 10,
                'internal_conflict': -5,  # CBT reduces internal conflict
                'growth_potential': 8
            },
            'ifs': {
                'emotional_regulation': 8,
                'internal_conflict': 15,  # IFS indicates internal parts conflict
                'growth_potential': 12,
                'relationship_sensitivity': 5
            },
            'attachment': {
                'relationship_sensitivity': 20,
                'emotional_regulation': -5,  # Attachment issues may reduce regulation
                'internal_conflict': 10,
                'growth_potential': 5
            },
            'jungian': {
                'cognitive_flexibility': 15,
                'growth_potential': 15,
                'internal_conflict': 8,  # Shadow work indicates internal conflict
                'emotional_regulation': 3
            },
            'narrative': {
                'cognitive_flexibility': 12,
                'growth_potential': 10,
                'internal_conflict': -8,  # Narrative therapy reduces conflict
                'emotional_regulation': 5
            }
        }
        
        # Apply framework-based adjustments
        for analysis in psychology_analyses:
            frameworks = analysis.get('frameworks', {})
            for framework_name, framework_data in frameworks.items():
                confidence = framework_data.get('confidence_score', 0.0)
                elements = framework_data.get('elements_detected', [])
                
                # Only apply if confidence is reasonable or elements detected
                if confidence > 0.5 or len(elements) >= 1:
                    if framework_name in framework_impacts:
                        impacts = framework_impacts[framework_name]
                        weight = min(confidence, 1.0)  # Cap at 1.0
                        
                        for score_type, impact in impacts.items():
                            adjustment = int(impact * weight)
                            scores[score_type] += adjustment
        
        # Ensure scores stay within 0-100 range
        for key in scores:
            scores[key] = max(0, min(100, scores[key]))
        
        return scores
    
    def _get_emotion_level(self, score: int) -> Dict[str, str]:
        """Get emotion level descriptions based on score"""
        if score >= 80:
            return {
                'identification': '准确',
                'regulation': '迅速',
                'risk': '稳定'
            }
        elif score >= 60:
            return {
                'identification': '清晰',
                'regulation': '较快',
                'risk': '适度'
            }
        elif score >= 40:
            return {
                'identification': '基础',
                'regulation': '一般',
                'risk': '敏感'
            }
        else:
            return {
                'identification': '初步',
                'regulation': '需要多些时间',
                'risk': '焦虑'
            }
    
    def _get_score_impact_analysis(self, score: int, score_type: str) -> str:
        """Generate impact analysis based on score"""
        
        impact_templates = {
            'emotional_regulation': {
                'high': "你的分值属于较高水平，显示你在情绪调节方面具有良好的基础能力。这表明你能够有效识别和管理情绪，在压力情况下保持相对稳定。",
                'medium': "你的分值属于中等偏高，仅反映你在情绪调节方面的当下状态，并不代表你的整体能力。这些分值显示你在情绪识别、调节和风险管理方面具备一定基础，但仍有提升空间，意味着通过练习和自我调节可以进一步增强情绪稳定性和应对能力。",
                'low': "你的分值显示在情绪调节方面可能面临一些挑战。这并不意味着能力不足，而是表明当前可能需要更多的支持和练习来发展情绪管理技能。"
            }
        }
        
        if score >= 70:
            level = 'high'
        elif score >= 40:
            level = 'medium'
        else:
            level = 'low'
        
        template = impact_templates.get(score_type, impact_templates['emotional_regulation'])
        return template[level]
    
    def _get_chinese_framework_interpretation(self, framework_name: str, data: Dict) -> str:
        """Get clinical interpretation for a framework in Chinese"""
        interpretations_chinese = {
            'ifs': "IFS模式的存在表明活跃的内在部分可能受益于自我领导发展和部分整合工作。",
            'cbt': "CBT模式表明认知和行为要素可能对结构化认知重构干预反应良好。",
            'jungian': "荣格元素表明丰富的无意识材料和原型内容可以支持个体化和个人成长。",
            'narrative': "叙事模式表明问题外化和发展首选身份故事的机会。",
            'attachment': "依恋模式表明植根于早期依恋经历的关系动态可能受益于依恋导向的干预。"
        }
        
        return interpretations_chinese.get(framework_name, "")
        """Get clinical interpretation for a framework in Chinese"""
        interpretations_chinese = {
            'ifs': "IFS模式的存在表明活跃的内在部分可能受益于自我领导发展和部分整合工作。",
            'cbt': "CBT模式表明认知和行为要素可能对结构化认知重构干预反应良好。",
            'jungian': "荣格元素表明丰富的无意识材料和原型内容可以支持个体化和个人成长。",
            'narrative': "叙事模式表明问题外化和发展首选身份故事的机会。",
            'attachment': "依恋模式表明植根于早期依恋经历的关系动态可能受益于依恋导向的干预。"
        }
        
        return interpretations_chinese.get(framework_name, "")


def generate_chinese_conversation_report(conversation_data: Dict, user_info: Optional[Dict] = None, output_dir: str = "reports") -> Optional[str]:
    """
    Convenience function to generate a Chinese report for a conversation.
    
    Args:
        conversation_data: Complete conversation data
        user_info: Optional user information
        output_dir: Output directory for reports
        
    Returns:
        Path to generated report file, or None if criteria not met
    """
    generator = ChineseTemplateReportGenerator(output_dir=output_dir)
    
    should_generate, reason = generator.should_generate_report(conversation_data)
    
    if not should_generate:
        logger.info(f"报告未生成: {reason}")
        return None
    
    logger.info(f"生成报告: {reason}")
    return generator.generate_report(conversation_data, user_info)