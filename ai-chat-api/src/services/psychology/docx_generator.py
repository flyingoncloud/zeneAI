"""
DOCX Report Generator
Converts markdown report to DOCX format with embedded charts
"""
import os
import re
from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from jinja2 import Template
import logging

logger = logging.getLogger(__name__)


class DocxReportGenerator:
    """Generate DOCX psychology reports from markdown template"""

    def __init__(self, template_path: str = None):
        """
        Initialize DOCX generator

        Args:
            template_path: Path to markdown template file
        """
        if template_path is None:
            # Default to the markdown template
            base_dir = Path(__file__).parent.parent.parent
            template_path = base_dir / "resources" / "ZeneMe - 内视觉察专业报告.md"

        self.template_path = Path(template_path)

        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {self.template_path}")

    def generate_docx(
        self,
        report_data: Dict[str, Any],
        output_path: str,
        charts_dir: str = None
    ) -> str:
        """
        Generate DOCX report from report data

        Args:
            report_data: Complete report data dictionary
            output_path: Path to save the DOCX file
            charts_dir: Directory containing generated chart images

        Returns:
            Path to generated DOCX file
        """
        try:
            # Step 1: Render markdown template with Jinja2
            rendered_markdown = self._render_template(report_data)

            # Step 2: Create DOCX document
            doc = Document()

            # Set document properties
            self._set_document_properties(doc)

            # Step 3: Parse markdown and convert to DOCX
            self._markdown_to_docx(doc, rendered_markdown, charts_dir)

            # Step 4: Save document
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))

            logger.info(f"DOCX report generated successfully: {output_path}")
            return str(output_path)

        except Exception as e:
            logger.error(f"Error generating DOCX report: {str(e)}")
            raise

    def _render_template(self, report_data: Dict[str, Any]) -> str:
        """Render markdown template with Jinja2"""
        with open(self.template_path, 'r', encoding='utf-8') as f:
            template_content = f.read()

        template = Template(template_content)
        return template.render(report_data)

    def _set_document_properties(self, doc: Document):
        """Set document-wide properties like fonts and styles"""
        # Set default font for the document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'SimSun'  # Chinese font
        font.size = Pt(11)

        # Set font for East Asian text
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    def _markdown_to_docx(self, doc: Document, markdown_text: str, charts_dir: str = None):
        """
        Convert markdown text to DOCX format

        Args:
            doc: Document object to add content to
            markdown_text: Rendered markdown text
            charts_dir: Directory containing chart images
        """
        lines = markdown_text.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            # Skip empty lines
            if not line:
                i += 1
                continue

            # Handle headers
            if line.startswith('#'):
                self._add_heading(doc, line)

            # Handle images
            elif line.startswith('!['):
                self._add_image(doc, line, charts_dir)

            # Handle tables
            elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                table_lines = [line]
                i += 1
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i].strip())
                    i += 1
                self._add_table(doc, table_lines)
                continue

            # Handle lists
            elif line.startswith('- ') or line.startswith('* ') or re.match(r'^\d+\.', line):
                self._add_list_item(doc, line)

            # Handle blockquotes
            elif line.startswith('>'):
                self._add_blockquote(doc, line)

            # Handle horizontal rules
            elif line.startswith('---'):
                self._add_horizontal_rule(doc)

            # Regular paragraph
            else:
                self._add_paragraph(doc, line)

            i += 1

    def _add_heading(self, doc: Document, line: str):
        """Add heading to document"""
        # Count # symbols to determine level
        level = 0
        while level < len(line) and line[level] == '#':
            level += 1

        # Extract heading text
        heading_text = line[level:].strip()

        # Add heading (DOCX supports levels 1-9)
        heading_level = min(level, 9)
        heading = doc.add_heading(heading_text, level=heading_level)

        # Style heading
        if level == 1:
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading.runs:
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue

    def _add_image(self, doc: Document, line: str, charts_dir: str = None):
        """Add image to document"""
        # Extract image path from markdown: ![alt](path)
        match = re.match(r'!\[.*?\]\((.*?)\)', line)
        if not match:
            return

        image_path = match.group(1)

        # If charts_dir provided, look for image there
        if charts_dir:
            image_path = os.path.join(charts_dir, os.path.basename(image_path))

        # Check if image exists
        if os.path.exists(image_path):
            try:
                # Add image with appropriate width
                doc.add_picture(image_path, width=Inches(5.5))
                # Center the image
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                logger.warning(f"Could not add image {image_path}: {str(e)}")
                # Add placeholder text
                doc.add_paragraph(f"[图片: {os.path.basename(image_path)}]")
        else:
            logger.warning(f"Image not found: {image_path}")
            doc.add_paragraph(f"[图片: {os.path.basename(image_path)}]")

    def _add_table(self, doc: Document, table_lines: list):
        """Add table to document"""
        # Parse table
        rows = []
        for line in table_lines:
            # Skip separator lines (e.g., | --- | --- |)
            if re.match(r'^\|[\s\-:]+\|', line):
                continue

            # Split by | and clean
            cells = [cell.strip() for cell in line.split('|')]
            # Remove empty first/last cells
            cells = [c for c in cells if c]
            if cells:
                rows.append(cells)

        if not rows:
            return

        # Create table
        num_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Light Grid Accent 1'

        # Fill table
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = table.rows[i].cells[j]
                    cell.text = cell_text

                    # Bold header row
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

    def _add_list_item(self, doc: Document, line: str):
        """Add list item to document"""
        # Remove list marker
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            style = 'List Bullet'
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            style = 'List Number'
        else:
            text = line
            style = 'List Bullet'

        # Add paragraph with list style
        p = doc.add_paragraph(text, style=style)

        # Apply inline formatting (bold, italic, etc.)
        self._apply_inline_formatting(p)

    def _add_blockquote(self, doc: Document, line: str):
        """Add blockquote to document"""
        text = line[1:].strip()  # Remove >
        p = doc.add_paragraph(text)

        # Style as quote
        p.paragraph_format.left_indent = Inches(0.5)
        for run in p.runs:
            run.font.italic = True
            run.font.color.rgb = RGBColor(96, 96, 96)

    def _add_horizontal_rule(self, doc: Document):
        """Add horizontal rule (line separator)"""
        # Use a simple line of dashes as horizontal rule
        p = doc.add_paragraph('─' * 80)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.color.rgb = RGBColor(192, 192, 192)
            run.font.size = Pt(8)

    def _add_paragraph(self, doc: Document, line: str):
        """Add regular paragraph to document"""
        # Skip if line is just formatting
        if not line or line in ['---', '***', '___']:
            return

        p = doc.add_paragraph(line)
        self._apply_inline_formatting(p)

    def _apply_inline_formatting(self, paragraph):
        """Apply inline markdown formatting (bold, italic, etc.)"""
        # This is a simplified version - for production, consider using a proper markdown parser
        text = paragraph.text

        # Clear existing runs
        paragraph.clear()

        # Handle bold (**text** or __text__)
        parts = re.split(r'(\*\*.*?\*\*|__.*?__)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith('__') and part.endswith('__'):
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
            else:
                # Handle italic (*text* or _text_)
                italic_parts = re.split(r'(\*.*?\*|_.*?_)', part)
                for ipart in italic_parts:
                    if ipart.startswith('*') and ipart.endswith('*') and not ipart.startswith('**'):
                        run = paragraph.add_run(ipart[1:-1])
                        run.font.italic = True
                    elif ipart.startswith('_') and ipart.endswith('_') and not ipart.startswith('__'):
                        run = paragraph.add_run(ipart[1:-1])
                        run.font.italic = True
                    else:
                        paragraph.add_run(ipart)


def generate_psychology_report_docx(
    report_data: Dict[str, Any],
    output_dir: str,
    report_id: int,
    charts_dir: str = None
) -> str:
    """
    Convenience function to generate psychology report DOCX

    Args:
        report_data: Complete report data dictionary
        output_dir: Directory to save the report
        report_id: Report ID for filename
        charts_dir: Directory containing chart images

    Returns:
        Path to generated DOCX file
    """
    output_path = os.path.join(output_dir, f"psychology_report_{report_id}.docx")

    generator = DocxReportGenerator()
    return generator.generate_docx(report_data, output_path, charts_dir)
